# _*_ coding: utf-8 _*_
# @Author  : XuTusheng
# @Time    : 2019/9/30
# @PROJECT : di5cheng-UT-Fleet-P
# @File    : ParseWord.py

import docx

from random import randint


class ParseWord:
    """
    解析word文件数据类
    """
    def __init__(self, filepath):
        """
        :param filepath: 文件路径
        """
        self.path = filepath

    def get_word(self):
        """
        获取word文件对象
        :return: 返回数据
        """
        try:
            doc = docx.Document(self.path)
            return doc
        except Exception as msg:
            print("异常消息-> {0}".format(msg))

    def fetchall_text(self):
        """获取word文件所有内容"""
        doc = self.get_word()
        fulltext = []
        for para in doc.paragraphs:
            fulltext.append(para.text)
        return '\n'.join(fulltext)

    def fetchone_text(self):
        """获取word文件某段内容"""
        doc = self.get_word()
        fulltext = []
        for para in doc.paragraphs:
            fulltext.append(para.text)
        num = randint(0, len(fulltext) - 1)
        return fulltext[num]
